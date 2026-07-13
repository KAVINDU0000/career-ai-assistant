"""
report_export.py
-------------------
Converts the generated career report (Markdown text) into downloadable
Word (.docx) and PDF files, with basic formatting: headings, bullet
points, bold text, and simple tables.

Both python-docx and fpdf2 are pure-Python with no system-level binary
dependencies (unlike e.g. weasyprint/pdfkit, which need external
libraries that have proven unreliable on constrained hosting like
Streamlit Community Cloud in this project already). That reliability is
the main reason they were chosen here.

This is a lightweight line-based Markdown parser, not a full CommonMark
implementation - it covers exactly the subset of Markdown this app's
agents actually produce: '##'/'###' headings, '-'/'*' bullet lists,
'**bold**' inline text, simple pipe tables, and plain paragraphs.
"""

import io
import re
from typing import List, Tuple

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

NAVY_RGB = (15, 23, 42)
BLUE_RGB = (37, 99, 235)

# fpdf2's base fonts (Helvetica, etc.) only support the Latin-1/WinAnsi
# character set. LLM-generated text very commonly contains "smart"
# punctuation - en/em dashes, curly quotes, ellipses - that falls outside
# that range and would otherwise crash PDF generation with
# FPDFUnicodeEncodingException. Map the common cases to safe ASCII
# equivalents, then use encode/decode as a final catch-all so ANY other
# unsupported character (emoji, foreign scripts, currency symbols, etc.)
# degrades to '?' instead of crashing the whole export.
_UNICODE_REPLACEMENTS = {
    "\u2013": "-",    # en dash
    "\u2014": "-",    # em dash
    "\u2018": "'",    # left single quote
    "\u2019": "'",    # right single quote
    "\u201c": '"',    # left double quote
    "\u201d": '"',    # right double quote
    "\u2026": "...",  # ellipsis
    "\u00a0": " ",    # non-breaking space
    "\u2022": "-",    # bullet
    "\u2192": "->",   # right arrow
    "\u2713": "v",    # check mark
    "\u2714": "v",    # heavy check mark
    "\u2715": "x",    # multiplication x
    "\u274c": "x",    # cross mark
}


def _sanitize_for_pdf(text: str) -> str:
    """
    Make text safe to pass to fpdf2's base (non-Unicode) fonts.

    Replaces common "smart" punctuation with ASCII equivalents, then
    encodes to Latin-1 with errors='replace' as a final safety net so any
    remaining unsupported character (rather than crashing the export)
    just becomes a '?' - this guarantees PDF generation can never crash
    on unexpected characters again, regardless of what an LLM produces.
    """
    for unicode_char, ascii_equivalent in _UNICODE_REPLACEMENTS.items():
        text = text.replace(unicode_char, ascii_equivalent)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _parse_inline_bold(text: str) -> List[Tuple[str, bool]]:
    """
    Split a line into (text, is_bold) chunks based on **bold** markers.

    Args:
        text: A single line of markdown text.

    Returns:
        List of (chunk_text, is_bold) tuples, in order.
    """
    parts = re.split(r"(\*\*.*?\*\*)", text)
    chunks = []
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            chunks.append((part[2:-2], True))
        else:
            chunks.append((part, False))
    return chunks


def _is_table_row(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def _is_table_separator(line: str) -> bool:
    """
    Detect a markdown table's separator row (e.g. '| --- | --- | --- |').

    Splits into individual cells first, then checks each cell contains
    only dashes/colons - a naive strip('|') on the whole line only removes
    the outermost pipe characters, not the ones between cells, so a
    multi-column separator row would otherwise fail this check.
    """
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return bool(cells) and all(c != "" and all(ch in "-:" for ch in c) for c in cells)


def _parse_table_row(line: str) -> List[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


# ---------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------

def markdown_to_docx_bytes(markdown_text: str, title: str = "Career Report") -> bytes:
    """
    Convert markdown report text into a formatted Word document.

    Args:
        markdown_text: The full report in Markdown.
        title: Document title, shown at the top.

    Returns:
        Raw .docx file bytes, ready for a download button.
    """
    doc = Document()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(*NAVY_RGB)

    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Tables
        if _is_table_row(stripped):
            table_lines = []
            while i < len(lines) and _is_table_row(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [_parse_table_row(l) for l in table_lines if not _is_table_separator(l)]
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for r, row_cells in enumerate(rows):
                    for c, cell_text in enumerate(row_cells):
                        if c < len(table.columns):
                            cell = table.cell(r, c)
                            cell.text = ""
                            p = cell.paragraphs[0]
                            for chunk, bold in _parse_inline_bold(cell_text):
                                run = p.add_run(chunk)
                                run.bold = bold or r == 0
                                run.font.size = Pt(9)
            doc.add_paragraph()
            continue

        # Headings
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(*BLUE_RGB)
            i += 1
            continue

        if stripped.startswith("## "):
            p = doc.add_paragraph()
            p.space_before = Pt(14)
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(*NAVY_RGB)
            i += 1
            continue

        # Bullets
        if stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("• "):
            bullet_text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            for chunk, bold in _parse_inline_bold(bullet_text):
                run = p.add_run(chunk)
                run.bold = bold
                run.font.size = Pt(10.5)
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        for chunk, bold in _parse_inline_bold(stripped):
            run = p.add_run(chunk)
            run.bold = bold
            run.font.size = Pt(10.5)
        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------

class _ReportPDF(FPDF):
    def header(self):
        pass

    def footer(self):
        self.set_y(-12)
        self.set_font("Helvetica", size=8)
        self.set_text_color(140, 140, 140)
        self.cell(0, 8, f"Page {self.page_no()}", align="C")


def markdown_to_pdf_bytes(markdown_text: str, title: str = "Career Report") -> bytes:
    """
    Convert markdown report text into a formatted PDF.

    Args:
        markdown_text: The full report in Markdown.
        title: Document title, shown at the top.

    Returns:
        Raw PDF file bytes, ready for a download button.
    """
    pdf = _ReportPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    pdf.set_margins(16, 16, 16)

    page_width = pdf.w - 2 * 16

    def write_line(w: float, h: float, text: str, **kwargs):
        """
        multi_cell() leaves the cursor's X position wherever the last
        rendered line ended (especially with align='center'), rather than
        resetting it to the left margin. If the next call also uses
        width=0 ("remaining space to the right margin from the CURRENT
        x"), that remaining space can shrink to near-zero and raise
        FPDFException. Resetting x to the left margin after every write
        avoids this entirely.
        """
        pdf.multi_cell(w, h, _sanitize_for_pdf(text), **kwargs)
        pdf.set_x(pdf.l_margin)

    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(*NAVY_RGB)
    write_line(0, 10, title, align="C")
    pdf.ln(4)
    pdf.set_x(pdf.l_margin)

    def write_inline(text: str, size: float, indent: float = 0):
        """
        Write a line with mixed bold/non-bold chunks (from **bold** markers).

        multi_cell() can't mix bold and regular text within one call, so
        inline formatting uses write() per chunk instead - write() still
        respects the page's right margin for wrapping, it just doesn't
        support a custom cell width the way multi_cell does, which is why
        this is only used for full-width-ish content like paragraphs and
        bullets rather than table cells.
        """
        pdf.set_x(pdf.l_margin + indent)
        pdf.set_text_color(30, 30, 30)
        for chunk, bold in _parse_inline_bold(text):
            pdf.set_font("Helvetica", "B" if bold else "", size=size)
            pdf.write(6, _sanitize_for_pdf(chunk))
        pdf.ln(7)
        pdf.set_x(pdf.l_margin)

    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Tables -> simple grid
        if _is_table_row(stripped):
            table_lines = []
            while i < len(lines) and _is_table_row(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [_parse_table_row(l) for l in table_lines if not _is_table_separator(l)]
            if rows:
                col_count = len(rows[0])
                col_width = page_width / col_count
                for r, row_cells in enumerate(rows):
                    pdf.set_font("Helvetica", "B" if r == 0 else "", size=8.5)
                    pdf.set_text_color(30, 30, 30)
                    pdf.set_x(pdf.l_margin)
                    y_before = pdf.get_y()
                    max_h = 6
                    for c, cell_text in enumerate(row_cells[:col_count]):
                        x = pdf.get_x()
                        y = pdf.get_y()
                        pdf.multi_cell(col_width, 5, _sanitize_for_pdf(cell_text), border=1)
                        max_h = max(max_h, pdf.get_y() - y)
                        pdf.set_xy(x + col_width, y)
                    pdf.set_y(y_before + max_h)
                    pdf.set_x(pdf.l_margin)
            pdf.ln(3)
            continue

        if stripped.startswith("### "):
            pdf.set_font("Helvetica", "B", 12.5)
            pdf.set_text_color(*BLUE_RGB)
            write_line(0, 7, stripped[4:].replace("**", ""))
            i += 1
            continue

        if stripped.startswith("## "):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 15)
            pdf.set_text_color(*NAVY_RGB)
            write_line(0, 8, stripped[3:].replace("**", ""))
            pdf.ln(1)
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("• "):
            bullet_text = stripped[2:].strip()
            write_inline(f"-  {bullet_text}", size=10, indent=4)
            i += 1
            continue

        write_inline(stripped, size=10.5, indent=0)
        i += 1

    return bytes(pdf.output())
